from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE = ROOT.parent
PROJECT_DOCX = ROOT / "docs" / "paper" / "菜市番茄生鲜度智能识别系统的设计与实现_杨佳浩_2023115226_最终含图版.docx"
ROOT_DOCX_CANDIDATES = [
    WORKSPACE / "菜市番茄生鲜度智能识别系统的设计与实现_杨佳浩_2023115226_最终含图版.docx",
    WORKSPACE / "菜市番茄生鲜度智能识别系统的设计与实现_杨佳浩_2023115226.docx",
]

CODE_FILES = [
    ("附录A 识别核心代码：src/tomato_freshness/recognizer.py", ROOT / "src" / "tomato_freshness" / "recognizer.py"),
    ("附录B Web服务代码：src/tomato_freshness/webapp.py", ROOT / "src" / "tomato_freshness" / "webapp.py"),
    ("附录C 命令行入口代码：src/tomato_freshness/cli.py", ROOT / "src" / "tomato_freshness" / "cli.py"),
    ("附录D 前端页面结构：web/index.html", ROOT / "web" / "index.html"),
    ("附录E 前端交互代码：web/static/app.js", ROOT / "web" / "static" / "app.js"),
]

CODE_FONT_NAME = "宋体"
CODE_FONT_SIZE_PT = 7.5  # Word 中文字号“六号”约为 7.5pt。


def set_run_font(run, bold: bool | None = None) -> None:
    """设置附录代码字体为宋体六号。"""
    run.font.name = CODE_FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CODE_FONT_NAME)
    run.font.size = Pt(CODE_FONT_SIZE_PT)
    if bold is not None:
        run.bold = bold


def style_code_paragraph(paragraph, bold: bool | None = None) -> None:
    """统一附录代码段落字号、行距和段距。"""
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for run in paragraph.runs:
        set_run_font(run, bold=bold)


def find_appendix_start(doc: Document) -> int:
    """定位正文中的附录标题，跳过目录。"""
    for index, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text in {"附 录", "附录"} and "\t" not in paragraph.text:
            return index
    raise ValueError("未找到附录标题")


def remove_appendix_content(doc: Document, start_index: int) -> None:
    """删除附录标题之后的旧代码内容。"""
    body = doc._body._element
    paragraphs = list(doc.paragraphs)
    for paragraph in paragraphs[start_index + 1 :]:
        body.remove(paragraph._p)


def add_code_line(doc: Document, text: str = "", bold: bool | None = None, align_center: bool = False) -> None:
    paragraph = doc.add_paragraph()
    if align_center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text if text else " ")
    set_run_font(run, bold=bold)
    style_code_paragraph(paragraph, bold=bold)


def read_code(path: Path) -> list[str]:
    if not path.exists():
        raise FileNotFoundError(path)
    return path.read_text(encoding="utf-8").splitlines()


def rebuild_appendix(docx_path: Path) -> None:
    doc = Document(docx_path)
    start_index = find_appendix_start(doc)
    appendix_title = doc.paragraphs[start_index]
    appendix_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in appendix_title.runs:
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

    remove_appendix_content(doc, start_index)

    add_code_line(doc, "以下为当前项目最新核心代码。代码已按功能模块组织，注释保留为中文，字体统一设置为宋体六号。", bold=False)
    add_code_line(doc, "运行入口：PYTHONPATH=src python -m tomato_freshness.webapp --port 8765 --seed-demo", bold=False)
    add_code_line(doc)

    for title, path in CODE_FILES:
        add_code_line(doc, title, bold=True)
        add_code_line(doc, "-" * min(len(title), 80))
        for line in read_code(path):
            # 保留空行，便于阅读代码结构。
            add_code_line(doc, line)
        add_code_line(doc)

    doc.save(docx_path)
    print(f"已更新附录代码并设置宋体六号：{docx_path}")


def main() -> None:
    for docx_path in [PROJECT_DOCX, *ROOT_DOCX_CANDIDATES]:
        if docx_path.exists():
            rebuild_appendix(docx_path)


if __name__ == "__main__":
    main()
