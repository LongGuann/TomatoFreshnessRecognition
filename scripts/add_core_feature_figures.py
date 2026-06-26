from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE = ROOT.parent
FIGURE_DIR = ROOT / "docs" / "figures" / "core_features"
PROJECT_DOCX = ROOT / "docs" / "paper" / "菜市番茄生鲜度智能识别系统的设计与实现_杨佳浩_2023115226_最终含图版.docx"
ROOT_DOCX = WORKSPACE / "菜市番茄生鲜度智能识别系统的设计与实现_杨佳浩_2023115226_最终含图版.docx"

FONT_CANDIDATES = [
    Path("/System/Library/Fonts/STHeiti Medium.ttc"),
    Path("/System/Library/Fonts/Hiragino Sans GB.ttc"),
    Path("/System/Library/Fonts/Supplemental/Songti.ttc"),
]


FEATURE_FIGURES = [
    {
        "heading": "5.2.1",
        "caption": "图5.1  用户登录与权限实现流程图",
        "filename": "fig5_1_login_permission.png",
        "title": "用户登录与权限实现",
        "steps": ["输入账号密码", "校验用户信息", "判断用户角色", "分配访问权限", "记录登录日志"],
        "note": "管理员拥有用户管理、统计分析和日志管理权限；检测员仅开放检测与查询功能。",
        "color": "#2563EB",
    },
    {
        "heading": "5.2.2",
        "caption": "图5.2  图像采集与预处理实现流程图",
        "filename": "fig5_2_image_preprocess.png",
        "title": "图像采集与预处理实现",
        "steps": ["摄像头拍摄/本地上传", "读取番茄图像", "质量检测", "尺寸归一化", "降噪与颜色转换", "输出标准化图像"],
        "note": "质量检测会拦截模糊、过暗、过曝或无明显番茄区域的无效图片。",
        "color": "#16A34A",
    },
    {
        "heading": "5.2.3",
        "caption": "图5.3  番茄生鲜度智能识别实现流程图",
        "filename": "fig5_3_freshness_recognition.png",
        "title": "番茄生鲜度智能识别实现",
        "steps": ["输入标准化图像", "提取颜色特征", "计算暗斑与纹理", "生成新鲜度得分", "映射四级结果", "输出置信度"],
        "note": "当前版本采用可解释启发式规则，后续可替换为 CNN 模型推理接口。",
        "color": "#DC2626",
    },
    {
        "heading": "5.2.4",
        "caption": "图5.4  数据存储与溯源实现流程图",
        "filename": "fig5_4_record_trace.png",
        "title": "数据存储与溯源实现",
        "steps": ["接收识别结果", "保存图像路径", "写入检测记录", "追加 JSONL/数据库", "按条件查询", "导出检测明细"],
        "note": "记录包含图像路径、等级、置信度、异常状态和检测时间，便于后续追溯。",
        "color": "#7C3AED",
    },
    {
        "heading": "5.2.5",
        "caption": "图5.5  数据可视化与预警实现流程图",
        "filename": "fig5_5_visual_warning.png",
        "title": "数据可视化与预警实现",
        "steps": ["读取检测记录", "统计等级分布", "计算合格率", "生成图表", "识别异常等级", "弹出预警提示"],
        "note": "轻微变质和严重变质会触发异常预警，并进入异常台账统计。",
        "color": "#F59E0B",
    },
]


def font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    for path in FONT_CANDIDATES:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def text_size(draw: ImageDraw.ImageDraw, text: str, text_font: ImageFont.ImageFont) -> tuple[int, int]:
    box = draw.textbbox((0, 0), text, font=text_font)
    return box[2] - box[0], box[3] - box[1]


def center_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, text_font: ImageFont.ImageFont) -> None:
    x1, y1, x2, y2 = box
    width, height = text_size(draw, text, text_font)
    draw.text((x1 + (x2 - x1 - width) / 2, y1 + (y2 - y1 - height) / 2), text, fill="#111827", font=text_font)


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: str = "#334155") -> None:
    draw.line((start, end), fill=fill, width=4)
    x1, y1 = start
    x2, y2 = end
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    size = 14
    draw.polygon(
        [
            (x2, y2),
            (x2 - ux * size + px * size * 0.55, y2 - uy * size + py * size * 0.55),
            (x2 - ux * size - px * size * 0.55, y2 - uy * size - py * size * 0.55),
        ],
        fill=fill,
    )


def draw_feature_figure(config: dict[str, object]) -> Path:
    """生成一个核心功能实现流程图。"""
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    path = FIGURE_DIR / str(config["filename"])
    image = Image.new("RGB", (1500, 860), "white")
    draw = ImageDraw.Draw(image)

    title_font = font(42)
    step_font = font(26)
    note_font = font(25)
    accent = str(config["color"])

    draw.rounded_rectangle((35, 35, 1465, 825), radius=26, outline="#CBD5E1", width=3)
    center_text(draw, (260, 60, 1240, 125), str(config["title"]), title_font)

    steps = list(config["steps"])  # type: ignore[arg-type]
    start_x, start_y = 90, 230
    box_w, box_h = 330, 92
    gap_x, gap_y = 70, 110

    positions: list[tuple[int, int, int, int]] = []
    for index, step in enumerate(steps):
        row = index // 3
        col = index % 3
        x1 = start_x + col * (box_w + gap_x)
        y1 = start_y + row * (box_h + gap_y)
        positions.append((x1, y1, x1 + box_w, y1 + box_h))

    for index, (step, box) in enumerate(zip(steps, positions)):
        draw.rounded_rectangle(box, radius=18, fill="#F8FAFC", outline=accent, width=4)
        center_text(draw, box, f"{index + 1}. {step}", step_font)

    for index in range(len(positions) - 1):
        current = positions[index]
        nxt = positions[index + 1]
        if index == 2:
            arrow(draw, ((current[0] + current[2]) // 2, current[3]), ((nxt[0] + nxt[2]) // 2, nxt[1]), accent)
        else:
            arrow(draw, (current[2], (current[1] + current[3]) // 2), (nxt[0], (nxt[1] + nxt[3]) // 2), accent)

    note_box = (165, 660, 1335, 750)
    draw.rounded_rectangle(note_box, radius=18, fill="#FFF7ED", outline="#FDBA74", width=3)
    center_text(draw, note_box, str(config["note"]), note_font)

    image.save(path)
    return path


def new_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def insert_picture_after(anchor: Paragraph, image_path: Path, caption: str) -> None:
    """在指定段落后插入图片和图注。"""
    caption_paragraph = new_paragraph_after(anchor)
    caption_paragraph.alignment = 1
    caption_paragraph.add_run(caption)

    picture_paragraph = new_paragraph_after(anchor)
    picture_paragraph.alignment = 1
    picture_paragraph.paragraph_format.keep_with_next = True
    picture_paragraph.paragraph_format.keep_together = True
    caption_paragraph.paragraph_format.keep_together = True
    picture_paragraph.add_run().add_picture(str(image_path), width=Inches(5.9))


def clean_existing_core_figures(doc: Document) -> None:
    """清理旧的图5.1-图5.5，避免重复插入。"""
    captions = {str(config["caption"]) for config in FEATURE_FIGURES}
    paragraphs = list(doc.paragraphs)
    for index, paragraph in enumerate(paragraphs):
        if paragraph.text.strip() in captions:
            if index > 0 and "w:drawing" in paragraphs[index - 1]._p.xml:
                remove_paragraph(paragraphs[index - 1])
            remove_paragraph(paragraph)


def find_anchor_after_heading(doc: Document, heading_prefix: str) -> Paragraph:
    """找到对应 5.2.x 标题后的说明段，作为插图锚点。"""
    paragraphs = doc.paragraphs
    for index, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        # 目录项通常包含制表符和页码，例如“5.2.1 ...\t13”，不能作为插图锚点。
        if text.startswith(heading_prefix) and "\t" not in paragraph.text:
            for candidate in paragraphs[index + 1 :]:
                candidate_text = candidate.text.strip()
                if candidate_text and not candidate_text.startswith("5.2.") and not candidate_text.startswith("5.3"):
                    return candidate
    raise ValueError(f"未找到小节锚点：{heading_prefix}")


def set_page_break_before_heading(doc: Document, heading_prefix: str) -> None:
    """给指定小节标题设置段前分页，避免标题落在页底与正文分离。"""
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith(heading_prefix) and "\t" not in paragraph.text:
            paragraph.paragraph_format.page_break_before = True
            p_pr = paragraph._p.get_or_add_pPr()
            if p_pr.find(qn("w:pageBreakBefore")) is None:
                p_pr.append(OxmlElement("w:pageBreakBefore"))
            return
    raise ValueError(f"未找到需要分页的小节标题：{heading_prefix}")


def update_document(docx_path: Path, figures: dict[str, Path]) -> None:
    doc = Document(docx_path)
    clean_existing_core_figures(doc)
    for config in FEATURE_FIGURES:
        anchor = find_anchor_after_heading(doc, str(config["heading"]))
        insert_picture_after(anchor, figures[str(config["heading"])], str(config["caption"]))
    set_page_break_before_heading(doc, "5.2.4")
    doc.save(docx_path)
    print(f"已更新论文：{docx_path}")


def main() -> None:
    figures = {str(config["heading"]): draw_feature_figure(config) for config in FEATURE_FIGURES}
    for docx_path in [PROJECT_DOCX, ROOT_DOCX]:
        if docx_path.exists():
            update_document(docx_path, figures)


if __name__ == "__main__":
    main()
