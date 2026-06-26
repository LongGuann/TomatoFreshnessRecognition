from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE = ROOT.parent
PROJECT_DOCX = ROOT / "docs" / "paper" / "菜市番茄生鲜度智能识别系统的设计与实现_杨佳浩_2023115226_最终含图版.docx"
ROOT_DOCX_CANDIDATES = [
    WORKSPACE / "菜市番茄生鲜度智能识别系统的设计与实现_杨佳浩_2023115226_最终含图版.docx",
    WORKSPACE / "菜市番茄生鲜度智能识别系统的设计与实现_杨佳浩_2023115226.docx",
]
SCREENSHOT_DIR = ROOT / "docs" / "screenshots" / "system_tests"


SCREENSHOTS = {
    "login": SCREENSHOT_DIR / "fig5_6_login_permission_real.png",
    "detect": SCREENSHOT_DIR / "fig5_7_detection_result_real.png",
    "preprocess": SCREENSHOT_DIR / "fig6_2_preprocess_quality_real.png",
    "records": SCREENSHOT_DIR / "fig6_4_record_query_real.png",
    "alerts": SCREENSHOT_DIR / "fig6_5_abnormal_warning_real.png",
    "stats": SCREENSHOT_DIR / "fig6_6_statistics_real.png",
    "tests": SCREENSHOT_DIR / "fig6_7_function_test_table_real.png",
}


def set_run_font(run, size: int = 12, bold: bool | None = None, font_name: str = "宋体") -> None:
    """统一新增正文的中英文字体，避免与原论文格式脱节。"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def style_paragraph(paragraph, size: int = 12, bold: bool | None = None) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold)


def find_body_paragraph(doc: Document, prefix: str):
    """查找正文标题，跳过目录中带制表符的条目。"""
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith(prefix) and "\t" not in paragraph.text:
            return paragraph
    raise ValueError(f"未找到正文标题：{prefix}")


def remove_between(doc: Document, start_prefix: str, end_prefix: str):
    """删除两个正文标题之间的所有段落和表格。"""
    start = find_body_paragraph(doc, start_prefix)
    end = find_body_paragraph(doc, end_prefix)
    body = doc._body._element
    children = list(body)
    start_index = children.index(start._p)
    end_index = children.index(end._p)
    for element in children[start_index:end_index]:
        body.remove(element)
    return end


def add_paragraph(
    anchor,
    text: str = "",
    style: str | None = None,
    size: int = 12,
    bold: bool | None = None,
    page_break_before: bool = False,
):
    paragraph = anchor.insert_paragraph_before(text, style=style)
    paragraph.paragraph_format.page_break_before = page_break_before
    style_paragraph(paragraph, size=size, bold=bold)
    return paragraph


def add_picture(anchor, image_path: Path, caption: str, width: float = 5.75) -> None:
    """插入系统真实运行截图和居中图注。"""
    if not image_path.exists():
        raise FileNotFoundError(f"截图不存在：{image_path}")

    picture = anchor.insert_paragraph_before()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.paragraph_format.keep_with_next = True
    picture.add_run().add_picture(str(image_path), width=Inches(width))

    caption_paragraph = anchor.insert_paragraph_before(caption)
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.keep_together = True
    style_paragraph(caption_paragraph, size=10)


def set_table_borders(table) -> None:
    """为模板默认无边框表格补充细边框。"""
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "B7C7BD")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def add_table(anchor, headers: list[str], rows: list[list[str]], col_widths: list[int] | None = None) -> None:
    """插入论文测试表，使用 Word 表格而非图片，便于后续修改。"""
    doc = anchor._parent
    table = doc.add_table(rows=1, cols=len(headers), width=Inches(6.0))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value

    set_table_borders(table)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            if col_widths and index < len(col_widths):
                set_cell_width(cell, col_widths[index])
            for paragraph in cell.paragraphs:
                style_paragraph(paragraph, size=9)

    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            style_paragraph(paragraph, size=9, bold=True)

    anchor._p.addprevious(table._tbl)


def rebuild_system_sections(docx_path: Path) -> None:
    doc = Document(docx_path)
    anchor = remove_between(doc, "5.3", "7 ")

    add_paragraph(anchor, "5.3 系统界面实现", style="Heading 2")
    add_paragraph(
        anchor,
        "系统界面采用本地 Web 管理控制台实现，前端页面位于 web/ 目录，后端服务由 tomato_freshness.webapp 提供。"
        "页面围绕菜市场番茄质检人员的日常操作设计，采用左侧模块导航和右侧工作区布局，包含用户权限、图像检测、"
        "记录查询、异常预警、统计分析和功能测试看板。系统截图均由本项目在本地浏览器真实运行后截取。",
    )
    add_picture(anchor, SCREENSHOTS["login"], "图5.6  用户登录与权限界面")
    add_picture(anchor, SCREENSHOTS["detect"], "图5.7  番茄生鲜度检测界面")
    add_picture(anchor, SCREENSHOTS["records"], "图5.8  检测记录查询界面")
    add_picture(anchor, SCREENSHOTS["stats"], "图5.9  数据统计分析界面")

    add_paragraph(anchor, "6 系统测试", style="Heading 1", page_break_before=True)
    add_paragraph(anchor, "6.1 测试目的", style="Heading 2")
    add_paragraph(
        anchor,
        "系统测试用于验证番茄生鲜度识别系统是否满足课程设计需求，重点检查登录权限、图像预处理、"
        "智能识别、记录溯源、异常预警和统计分析等核心功能是否能够在真实运行页面中稳定执行。",
    )

    add_paragraph(anchor, "6.2 测试环境", style="Heading 2")
    add_table(
        anchor,
        ["测试项", "测试环境"],
        [
            ["硬件环境", "MacBook 本地开发环境，普通台式机或触摸屏一体机均可访问 Web 页面"],
            ["软件环境", "Python 3.12、Pillow、NumPy、主流浏览器 Chrome/Edge"],
            ["运行方式", "PYTHONPATH=src python -m tomato_freshness.webapp --port 8765 --seed-demo"],
            ["测试数据", "examples/ 演示图片、datasets/test_set/ 合成测试集 48 张、outputs/web_records.jsonl 检测记录"],
            ["截图方式", "使用 Playwright 打开 http://127.0.0.1:8765 后截取真实系统页面"],
        ],
        col_widths=[1800, 7200],
    )

    add_paragraph(anchor, "6.3 功能测试", style="Heading 2", page_break_before=True)
    add_paragraph(anchor, "针对系统核心功能设计测试用例，测试结果如表6.1所示。")
    add_paragraph(anchor, "表6.1  功能测试用例表", size=10)
    add_table(
        anchor,
        ["测试项", "输入数据", "预期结果", "实际结果", "结论"],
        [
            ["登录权限测试", "admin/admin123、inspector/inspect123", "管理员和检测员显示不同权限", "权限标签正常加载", "通过"],
            ["图像预处理测试", "正常番茄图、模糊图", "系统完成质量检测并拦截无效图像", "模糊图返回无效图像", "通过"],
            ["识别功能测试", "优质、新鲜合格、轻微变质、严重变质样本", "返回等级、置信度和预警信息", "页面展示识别结果和特征值", "通过"],
            ["记录查询测试", "按等级和异常状态筛选", "返回匹配检测记录", "记录列表正常刷新", "通过"],
            ["异常预警测试", "轻微变质、严重变质样本", "触发异常预警并进入台账", "预警区显示处理建议", "通过"],
            ["统计分析测试", "多条检测记录", "生成总量、合格率、异常占比和等级分布", "统计卡片和柱状图正常显示", "通过"],
        ],
        col_widths=[1500, 2200, 2400, 2100, 800],
    )

    function_screens = [
        ("6.3.1 登录权限测试", "系统使用内置演示账号进行权限验证。管理员登录后可查看检测、查询、统计、用户管理和日志管理权限；检测员账号只开放检测与查询权限。", SCREENSHOTS["login"], "图6.1  登录权限测试截图"),
        ("6.3.2 图像预处理测试", "选择模糊样本后，系统完成尺寸归一化、清晰度检测和质量判断，识别结果被标记为无效图像，提示重新采集番茄图片。", SCREENSHOTS["preprocess"], "图6.2  图像预处理质量拦截图"),
        ("6.3.3 识别功能测试", "选择优质番茄样本并点击开始识别后，系统返回新鲜度等级、置信度、新鲜度得分和关键图像特征。", SCREENSHOTS["detect"], "图6.3  生鲜度识别功能测试截图"),
        ("6.3.4 记录查询测试", "系统将每次检测的图片、等级、置信度、检测人员和时间写入记录文件，并支持按等级和异常状态筛选。", SCREENSHOTS["records"], "图6.4  检测记录查询测试截图"),
        ("6.3.5 异常预警测试", "当识别结果为轻微变质、严重变质或无效图像时，系统在异常预警区展示红色提示和处理建议。", SCREENSHOTS["alerts"], "图6.5  异常预警测试截图"),
        ("6.3.6 统计分析测试", "系统基于检测记录生成检测总量、有效检测数、合格率、异常占比和等级分布柱状图。", SCREENSHOTS["stats"], "图6.6  统计分析测试截图"),
        ("6.3.7 测试结果汇总", "系统功能测试页面汇总展示各项测试输入、预期结果和测试结论，便于课程验收时核对。", SCREENSHOTS["tests"], "图6.7  系统功能测试汇总截图"),
    ]
    for heading, paragraph, screenshot, caption in function_screens:
        add_paragraph(anchor, heading, style="Heading 3")
        add_paragraph(anchor, paragraph)
        add_picture(anchor, screenshot, caption)

    add_paragraph(anchor, "6.4 性能测试", style="Heading 2")
    add_paragraph(
        anchor,
        "在本地 Web 服务启动后，对 /api/recognize 识别接口连续请求 10 次，测试图片为 examples/tomato_excellent.png。"
        "实测平均响应时间为 8.53ms，最大响应时间为 28.87ms，最小响应时间为 5.98ms。当前系统使用启发式规则识别，"
        "不加载大型深度学习模型，因此响应速度能够满足课程演示和小规模菜市场质检场景需求。",
    )
    add_table(
        anchor,
        ["测试接口", "请求次数", "平均响应", "最大响应", "测试结论"],
        [["/api/recognize", "10次", "8.53ms", "28.87ms", "通过"]],
        col_widths=[2400, 1400, 1600, 1600, 1200],
    )

    add_paragraph(anchor, "6.5 稳定性测试", style="Heading 2")
    add_paragraph(
        anchor,
        "系统连续完成登录、样例切换、识别、记录刷新、异常预警和统计展示流程，未出现页面空白、接口报错或服务中断。"
        "检测记录以 JSON Lines 形式追加保存，即使单次识别结果为无效图像，也会保留质量拦截信息，便于后续追溯。",
    )

    add_paragraph(anchor, "6.6 兼容性测试", style="Heading 2")
    add_paragraph(
        anchor,
        "前端页面采用标准 HTML、CSS 和 JavaScript 实现，不依赖复杂前端框架；截图测试在 Chromium 浏览器中完成。"
        "页面使用响应式布局，普通电脑、触摸屏一体机和较窄屏幕均可访问主要功能区。",
    )

    add_paragraph(anchor, "6.7 测试总结", style="Heading 2")
    add_paragraph(
        anchor,
        "测试结果表明，系统已完成课程设计所需的核心功能闭环：用户登录后可进行番茄图像检测，系统能够完成预处理、"
        "生鲜度等级识别、异常预警、记录溯源和统计分析。当前版本主要用于功能演示，识别算法采用可解释启发式规则，"
        "后续可在保持接口不变的基础上接入 CNN 模型以提升真实数据集上的识别准确率。",
    )

    doc.save(docx_path)
    print(f"已更新系统界面与系统测试章节：{docx_path}")


def main() -> None:
    for docx_path in [PROJECT_DOCX, *ROOT_DOCX_CANDIDATES]:
        if docx_path.exists():
            rebuild_system_sections(docx_path)


if __name__ == "__main__":
    main()
